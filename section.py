import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Список секций из информационного письма
SECTIONS = [
    "1. Интеллектуальные системы управления и принятия решений",
    "2. Информационные технологии при построении средств обработки информации",
    "3. Современные технологии при построении средств автоматизации",
    "4. Искусственный интеллект в управлении, автоматике и обработке информации",
    "5. Современные технологии в проектировании авиакосмических систем",
    "6. Информационные технологии в измерительных и вычислительных системах и сетях",
    "7. Современные технологии в задачах навигации и ориентации",
    "8. Новые технологии в прикладной и гуманитарной сфере"
]
import re  # Добавьте импорт регулярных выражений

def clean_text_docx(text):
    """Удаляет лишние пробелы и артефакты"""
    if not isinstance(text, str):
        text = str(text)
        
    # Удаляем специальные символы
    text = text.replace('\r', ' ').replace('\n', ' ').replace('_x000D_', ' ')
    
    # Удаляем лишние пробелы:
    text = re.sub(r'\s+', ' ', text)       # Множественные пробелы → один
    text = re.sub(r'\s([.,;!?])', r'\1', text)  # Убираем пробелы перед знаками препинания
    text = re.sub(r'([\(])\s', r'\1', text)    # Убираем пробелы после открывающей скобки
    text = re.sub(r'\s([\)])', r'\1', text)    # Убираем пробелы перед закрывающей скобкой
    
    return text.strip().strip('"')

def clean_text(text):
    """Очистка текста от артефактов"""
    if not isinstance(text, str):
        text = str(text)
    return text.replace('\r', ' ').replace('\n', ' ').replace('_x000D_', ' ').strip().strip('"')

# Определение секции по столбцу Menu 303
def detect_section(menu_303):
    try:
        # Извлекаем номер секции из начала значения Menu 303
        section_num = int(menu_303.split('.')[0])
        return SECTIONS[section_num - 1]
    except (ValueError, IndexError):
        return "Без секции"

# Загрузка данных
# В sheet_name заносим нужный лист в xlsx (обычно это Лист 1)
data = pd.read_excel("file.xlsx", sheet_name="cfdb7-2025-04-02 (1)", header=0)

# Создание единого документа
doc = Document()

# Группировка по секциям через столбец Menu 303
grouped = data.groupby(lambda idx: detect_section(clean_text(data.loc[idx, "Menu 303"])))

for section_title, group in grouped:
    # Добавляем заголовок секции
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(section_title.upper())
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = True
    
    # Пропуск после заголовка секции
    doc.add_paragraph().add_run().font.size = Pt(10)
    
    # Добавляем доклады секции
    for _, row in group.iterrows():
        # 1. Название доклада (рус)
        title_ru = clean_text(row["Name Title"]).upper()
        p = doc.add_paragraph(title_ru)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.bold = True
        
        
        # 2. Пропуск 6пт
        doc.add_paragraph().add_run().font.size = Pt(6)
        
        # 3. Название (англ)
        title_en = clean_text(row["English Name"])
        p = doc.add_paragraph(title_en)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.italic = True
        
        # 4. Пропуск 6пт
        doc.add_paragraph().add_run().font.size = Pt(6)
        
        # 5. Авторы
        authors = [clean_text(a) for a in str(row["Authors"]).split(";") if a.strip()]
        p = doc.add_paragraph(", ".join(authors))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.italic = True
        
        # 6. Пропуск 6пт
        doc.add_paragraph().add_run().font.size = Pt(6)
        
        # 7. Организация
        org = clean_text(row["Organisation Name"])
        p = doc.add_paragraph(org)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        
        # 8. Кафедра и руководитель
        department = clean_text(row.get("Department", ""))
        if department:
            p = doc.add_paragraph(department)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.italic = True
        
        # 9. Пропуск 10пт
        doc.add_paragraph().add_run().font.size = Pt(10)
        
        # 10. Аннотация (заголовок)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Аннотация")
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.bold = True
        
        # 11. Текст аннотации
        annotation = clean_text(row["Annotation"])
        p = doc.add_paragraph(annotation)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run()
        run.font.name = "Arial"
        run.font.size = Pt(8)
        
        # 12. Пропуск 10пт
        doc.add_paragraph().add_run().font.size = Pt(10)
        
        # 13. Текст доклада
        report_text = clean_text_docx(row["Report Text"])
        paragraphs = report_text.split('\n')

        for para_text in paragraphs:
            if para_text.strip():  # Пропускаем пустые строки
                p = doc.add_paragraph(para_text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Cm(1.25)  # Красная строка 1.25 см
                
                # Форматируем каждый абзац
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)
        
        # 14. Пропуск 6пт
        doc.add_paragraph().add_run().font.size = Pt(6)
        
        # 15. Список литературы (заголовок)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Список литературы")
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.italic = True
        
        # 16. Список литературы
        references = [clean_text(ref).lstrip(". ") for ref in str(row["Book List"]).split(";") if ref.strip()]
        for ref in references:
            p = doc.add_paragraph(ref)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run()
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.italic = True
        
        # Разделитель между докладами
        doc.add_paragraph().add_run().font.size = Pt(10)

# Сохранение общего документа
doc.save("conference_sections.docx")
print("Документ с секциями успешно создан!")