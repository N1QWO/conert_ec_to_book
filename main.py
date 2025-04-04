import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_custom_spacing(paragraph, before=0, after=0):
    """Установка интервалов до и после абзаца"""
    p = paragraph._element
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    p.get_or_add_pPr().append(spacing)

def setup_page(doc):
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # Высота A4
    section.page_width = Cm(21)     # Ширина A4
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.5)  # Расстояние до верхнего колонтитула
    section.footer_distance = Cm(0)     # Расстояние до нижнего колонтитула
    
    # Отключаем связь колонтитулов с предыдущими секциями
    if section.header is not None:
        section.header.is_linked_to_previous = False
    if section.footer is not None:
        section.footer.is_linked_to_previous = False

def clean_text(text):
    """Очистка текста от артефактов"""
    if not isinstance(text, str):
        text = str(text)
    return text.replace('\r', ' ').replace('\n', ' ').replace('_x000D_', ' ').strip().strip('"')

# Загрузка данных
data = pd.read_excel("file.xlsx", sheet_name="cfdb7-2025-04-02 (1)", header=0)

for index, row in data.iterrows():
    doc = Document()
    setup_page(doc)
    
    # 1. Название доклада (рус)
    title_ru = clean_text(row["Name Title"]).upper()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title_ru)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.bold = True
    run.font.no_proof = True  # Отключаем проверку правописания для прописных
    p.paragraph_format.line_spacing = 1.0  # Без интервала
    
    # 2. Пропуск 6 пт
    p = doc.add_paragraph()
    p.add_run().font.size = Pt(6)
    
    # 3. Название (англ)
    title_en = clean_text(row["English Name"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title_en)
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.italic = True
    
    # 4. Пропуск 6 пт
    p = doc.add_paragraph()
    p.add_run().font.size = Pt(6)
    
    # 5. Авторы
    authors = [clean_text(a) for a in str(row["Authors"]).split(";") if a.strip()]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, author in enumerate(authors):
        run = p.add_run(author)
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.italic = True
        if i < len(authors)-1:
            p.add_run(", ")
    
    # 6. Пропуск 6 пт
    p = doc.add_paragraph()
    p.add_run().font.size = Pt(6)
    
    # 7. Организация
    org = clean_text(row["Organisation Name"])
    p = doc.add_paragraph(org)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    
    # 8. Кафедра и руководитель
    department = clean_text(row.get("Department", ""))
    if department:
        p = doc.add_paragraph(department)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.italic = True
    
    # 9. Пропуск 10 пт
    p = doc.add_paragraph()
    p.add_run().font.size = Pt(10)
    
    # 10. Аннотация (заголовок)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Аннотация")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.bold = True
    
    # 11. Текст аннотации
    annotation = clean_text(row["Annotation"])
    p = doc.add_paragraph(annotation)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
    
    # 12. Пропуск 10 пт
    p = doc.add_paragraph()
    p.add_run().font.size = Pt(10)
    
    # 13. Текст доклада
    report_text = clean_text(row["Report Text"])
    p = doc.add_paragraph(report_text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
    
    # 14. Пропуск 6 пт
    p = doc.add_paragraph()
    p.add_run().font.size = Pt(6)
    
    # 15. Список литературы (заголовок)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Список литературы")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    run.italic = True
    
    # 16. Список литературы
    references = [clean_text(ref) for ref in str(row["Book List"]).split(";") if ref.strip()]
    for ref in references:
        p = doc.add_paragraph(ref.lstrip(" ."))
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.italic = True
    
    # Сохранение документа
    filename = f"thesis_{index+1}.docx"
    doc.save(filename)
    print(f"Создан файл: {filename}")